import os
import uuid
from pathlib import Path
from fastapi import UploadFile
from app.core.config import settings


class FileService:
    def save(self, file_content: bytes, filename: str, org_id: str, meeting_id: str) -> str:
        if settings.AWS_S3_BUCKET:
            return self._save_s3(file_content, filename, org_id, meeting_id)
        return self._save_local(file_content, filename, org_id, meeting_id)

    def _save_local(self, file_content: bytes, filename: str, org_id: str, meeting_id: str) -> str:
        dir_path = Path(settings.UPLOAD_DIR) / org_id / meeting_id
        dir_path.mkdir(parents=True, exist_ok=True)
        file_path = dir_path / filename
        with open(file_path, "wb") as f:
            f.write(file_content)
        return str(file_path)

    def _save_s3(self, file_content: bytes, filename: str, org_id: str, meeting_id: str) -> str:
        import boto3
        client = boto3.client(
            "s3",
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            endpoint_url=settings.AWS_ENDPOINT_URL,
        )
        key = f"{org_id}/{meeting_id}/{filename}"
        client.put_object(Bucket=settings.AWS_S3_BUCKET, Key=key, Body=file_content)
        return f"s3://{settings.AWS_S3_BUCKET}/{key}"

    def extract_text(self, path: str) -> str:
        p = Path(path)
        suffix = p.suffix.lower()
        if suffix in (".txt", ".md"):
            with open(p, "r", encoding="utf-8") as f:
                return f.read()
        elif suffix == ".docx":
            from docx import Document
            doc = Document(str(p))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")


file_service = FileService()
